"""
Builds maintenance/FINAL_REPORT.docx - a readable, image-embedded Word
version of FINAL_REPORT.md, for submission/sharing.

Run with: .venv/Scripts/python maintenance/build_final_report_docx.py
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x55, 0x55, 0x55)


def set_cell_shading(cell, color_hex):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def add_title_page(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Software Maintenance Simulation")
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = ACCENT

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("Final Report — Corrective, Adaptive, Preventive & Perfective Maintenance")
    run2.font.size = Pt(16)
    run2.font.color.rgb = GREY

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in [
        "Project: ZenTask Pro (to-do) — FastAPI + SQLite + vanilla-JS task manager",
        "Repository: github.com/NafNawal04/to-do",
        "Branches: fix/search-like-wildcard-escape · adapt/datetime-utcnow-upgrade ·",
        "preventive/app-js-fetch-dedup · perfective/get-tasks-query-indexing — all merged into main",
    ]:
        r = meta.add_run(line + "\n")
        r.font.size = Pt(11)
        r.font.color.rgb = GREY
    doc.add_page_break()


def add_toc(doc):
    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = "Right-click and choose 'Update Field' to generate the table of contents."
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    r_element.append(fldChar4)
    doc.add_page_break()


def add_para(doc, text, size=11, bold=False, italic=False, color=None, space_after=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_rich_para(doc, segments, size=11, space_after=8):
    """segments: list of (text, bold) tuples."""
    p = doc.add_paragraph()
    for text, bold in segments:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullets(doc, items, size=11):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        if isinstance(item, tuple):
            label, rest = item
            r1 = p.add_run(label)
            r1.font.bold = True
            r1.font.size = Pt(size)
            r2 = p.add_run(rest)
            r2.font.size = Pt(size)
        else:
            r = p.add_run(item)
            r.font.size = Pt(size)


def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F2F2")
    cell.text = ""
    lines = text.strip("\n").split("\n")
    p0 = cell.paragraphs[0]
    for i, line in enumerate(lines):
        if i == 0:
            r = p0.add_run(line)
        else:
            p0.add_run().add_break()
            r = p0.add_run(line)
        r.font.name = 'Consolas'
        r.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_image(doc, rel_path, caption, max_width=6.3):
    full_path = os.path.join(HERE, rel_path)
    if not os.path.exists(full_path):
        add_para(doc, f"[Screenshot missing: {rel_path}]", italic=True, color=RGBColor(0xAA, 0x00, 0x00))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(full_path, width=Inches(max_width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = GREY
    cap.paragraph_format.space_after = Pt(12)


def add_cell_section(doc, num, name, tools, why, steps, image_info, interpretation):
    doc.add_heading(f"{num} {name}", level=3)
    add_rich_para(doc, [("Tools used: ", True), (tools, False)])
    add_para(doc, why, space_after=6)
    if steps:
        add_para(doc, "Command / steps:", bold=True, space_after=2)
        add_code_block(doc, steps)
    if image_info:
        for rel_path, caption in image_info:
            add_image(doc, rel_path, caption)
    add_para(doc, "Interpretation:", bold=True, space_after=2)
    add_para(doc, interpretation, space_after=14)


# ============================================================
doc = Document()

# Base font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

for section in doc.sections:
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

add_title_page(doc)
add_toc(doc)

# ---------------- Introduction ----------------
doc.add_heading("1. Overview", level=1)
add_para(doc,
    "This report documents a complete software maintenance simulation carried out against a real, "
    "working project — ZenTask Pro, a FastAPI + SQLite backend with a vanilla-JavaScript frontend. "
    "The exercise covers all four recognized categories of software maintenance — Corrective, Adaptive, "
    "Preventive, and Perfective — applied to the five standard maintenance tasks: Program Comprehension, "
    "Change Management, Impact Analysis, Reverse Engineering, and Refactoring. That is 20 cells in total, "
    "each backed by a real tool run against real code, not a hypothetical walkthrough.")
add_para(doc,
    "Every bug, deprecation, duplication cluster, and performance bottleneck described in this report is "
    "genuine — found by actually reading the code and running the tools against it, not invented for the "
    "exercise. Where a number is reported (timings, clone counts, SonarQube metrics), it is a number a "
    "tool actually produced, including the unflattering ones — for example, the perfective-maintenance "
    "fix's modest 19% speedup is reported honestly rather than inflated.")
add_para(doc,
    "Four branches were created, one per maintenance type, and all four were eventually merged back into "
    "main: fix/search-like-wildcard-escape (corrective), adapt/datetime-utcnow-upgrade (adaptive), "
    "preventive/app-js-fetch-dedup (preventive), and perfective/get-tasks-query-indexing (perfective).")

# ---------------- Why these tools ----------------
doc.add_heading("2. Why These Tools, and Why Not Others", level=1)
add_para(doc,
    "The available toolset for this exercise was: Graphviz, AST Explorer, SonarQube, Joern, Neo4j, "
    "Loguru, PySnooper, Viztracer, Snakeviz, cProfile, and CCFinderSW. Two commonly-used maintenance "
    "tools were deliberately excluded from this simulation:")
add_bullets(doc, [
    ("Doxygen — ", "its role (auto-generated structural documentation) is covered instead by AST Explorer "
     "plus a manually-written explanation report for every comprehension cell. A human-reasoned "
     "“here is what this code does and why it matters” writeup is a more honest and more useful "
     "comprehension artifact for a maintenance exercise than auto-generated docs would be."),
    ("IDA Pro — ", "built for compiled-binary disassembly. This is a pure Python/JavaScript source "
     "project with no compiled binary to disassemble, so including IDA Pro would add nothing real to "
     "the exercise."),
])
add_para(doc,
    "Beyond that, each of the 20 cells below used the specific tool(s) the assignment mapped to that "
    "maintenance-type/task combination. The table below is the map; the per-cell sections that follow "
    "explain why that tool was the right one for that specific finding, not simply that the plan called "
    "for it.")

add_table(doc,
    ["Tool", "What it is for", "Used in"],
    [
        ["PySnooper", "Line-by-line runtime trace of a specific function call", "Corrective (comprehension, refactor verification)"],
        ["AST Explorer", "Visualizes a function's syntax tree to confirm structural claims", "All 4 types, Program Comprehension"],
        ["Git", "Branching, commit history, change documentation", "All 4 types, Change Management"],
        ["Joern", "Generates a Code Property Graph (CPG) and queries it (CPGQL) for callers, callees, call fan-out, and cross-file dependencies", "All 4 types, Impact Analysis"],
        ["Neo4j", "Visualizes the CPG-derived call/dependency graph interactively", "All 4 types, Impact Analysis"],
        ["Graphviz", "Renders a call graph / dependency graph as a static diagram", "Corrective & Adaptive (Reverse Engineering), Preventive (Comprehension)"],
        ["SonarQube Cloud", "Static analysis: bugs, code smells, complexity, duplication, coverage", "Corrective, Adaptive, Preventive (Refactoring); Preventive (Impact Analysis)"],
        ["CCFinderSW", "Token-based clone detection across a codebase", "Preventive (Reverse Engineering + Refactoring re-run)"],
        ["Viztracer", "Visual call-timeline trace of one function execution", "Perfective (Comprehension)"],
        ["cProfile", "Exact call-count / timing profile of a function", "Perfective (Reverse Engineering + Refactoring re-check)"],
        ["Snakeviz", "Interactive flame-graph visualization of a .prof file", "Perfective (Reverse Engineering + Refactoring re-check)"],
        ["Loguru", "Structured logging to verify behavior after a change", "Perfective (Refactoring verification)"],
    ])

# ---------------- Summary table ----------------
doc.add_heading("3. Summary Table", level=1)
add_table(doc,
    ["Maintenance Type", "Program Comprehension", "Change Management", "Impact Analysis", "Reverse Engineering", "Refactoring"],
    [
        ["1. Corrective", "PySnooper + AST Explorer + Report", "Git", "Joern + Neo4j", "Graphviz", "SonarQube + PySnooper"],
        ["2. Adaptive", "AST Explorer + Report", "Git", "Joern + Neo4j", "Graphviz", "SonarQube"],
        ["3. Preventive", "Graphviz + AST Explorer + Report", "Git", "SonarQube + Neo4j (via Joern)", "CCFinderSW", "SonarQube + CCFinderSW"],
        ["4. Perfective", "Viztracer + AST Explorer + Report", "Git", "Joern + Neo4j", "cProfile + Snakeviz", "Loguru + cProfile/Snakeviz"],
    ])
doc.add_page_break()

# ============================================================
# 1. CORRECTIVE
# ============================================================
doc.add_heading("4. Corrective Maintenance", level=1)
add_para(doc,
    "Scenario: crud.get_tasks() — the function behind GET /api/tasks?search=... — built a SQL ILIKE "
    "pattern from the raw search query parameter using plain string interpolation (f\"%{search}%\"), "
    "never escaping SQL LIKE wildcard characters (% and _). A user searching for a literal underscore or "
    "percent sign got the wildcard interpreted instead of matched literally — for example, searching for "
    "\"_\" returned every task instead of only the one whose title actually contained an underscore. This "
    "was found by tracing the function with PySnooper on that exact edge case.", space_after=4)
add_rich_para(doc, [("Branch: ", True), ("fix/search-like-wildcard-escape  ", False),
                     ("Fix commit: ", True), ("f4dc3cf  ", False),
                     ("Status: ", True), ("merged into main", False)], size=10, space_after=14)

add_cell_section(doc, "4.1", "Program Comprehension",
    "PySnooper + AST Explorer + written report",
    "PySnooper gives a runtime trace showing the actual (wrong) SQL pattern being built and the actual "
    "(wrong) result set, which is the fastest way to see a data bug happen line-by-line. AST Explorer then "
    "confirms structurally that no escaping step exists anywhere in the function, ruling out “maybe "
    "it's escaped somewhere I'm not seeing.”",
    "Seeded an isolated in-memory database with 4 sample tasks (only 1 containing a literal underscore),\n"
    "wrapped crud.get_tasks(search=\"_\") in @pysnooper.snoop().",
    [("corrective/01_program_comprehension/ast_explorer_get_tasks_before_fix.png",
      "Figure 1.1 — AST Explorer: get_tasks() before the fix, showing no escape/sanitize call in the search branch")],
    "The trace showed get_tasks() returning all 4 tasks instead of 1; the AST confirmed the if search: "
    "branch has exactly one ilike() call chain with no sanitizing call anywhere in it — the bug is "
    "structural, not a one-off typo.")

add_cell_section(doc, "4.2", "Change Management",
    "Git",
    "Git is the tool for exactly this job: branching, bug reporting, and commit history are inherently "
    "text/diff-based, and nothing else in the toolset performs this role.",
    "git checkout -b fix/search-like-wildcard-escape\n"
    "# ... wrote bug_report.md (what broke, how found, expected vs. actual)\n"
    "git commit -m \"fix: escape LIKE wildcard characters in task search\"   # f4dc3cf",
    None,
    "Standard, low-risk change tracking. The bug report and git log evidence together document what broke, "
    "how it was found, and exactly which commit fixed it.")

add_cell_section(doc, "4.3", "Impact Analysis",
    "Joern + Neo4j",
    "Joern's Code Property Graph (CPG) can answer “who calls this function, anywhere in the "
    "codebase” precisely, which matters before changing a function's behavior — the blast radius must "
    "be known before a fix can be called safe.",
    "joern-parse --language pythonsrc <backend source>\n"
    "cpg.method.name(\"get_tasks\").caller.fullName.foreach(println)",
    [("corrective/03_impact_analysis/neo4j_impact_analysis.png",
      "Figure 1.2 — Neo4j Browser: get_tasks() call graph")],
    "get_tasks() has exactly one caller in the whole codebase — main.read_tasks(). The fix is fully "
    "contained to one HTTP endpoint; nothing else needed re-checking.")

add_cell_section(doc, "4.4", "Reverse Engineering",
    "Graphviz",
    "Once Joern has already answered “who calls this,” a small, purpose-built call graph is the "
    "clearest way to show where a fault originates and how it propagates, without the overhead of "
    "visualizing the entire CPG.",
    None,
    [("corrective/04_reverse_engineering/bug_callgraph.png",
      "Figure 1.3 — Graphviz call graph tracing the fault from the HTTP request to the wrong SQL result")],
    "The diagram traces the fault from the HTTP request, through read_tasks -> get_tasks -> the two "
    "ilike() calls (the fault origin), down to the wrong SQL result — a one-glance picture of the whole "
    "bug's lifecycle.")

add_cell_section(doc, "4.5", "Refactoring",
    "SonarQube + PySnooper",
    "PySnooper re-confirms the specific bug is fixed on the same edge case; SonarQube checks that the fix "
    "did not introduce new, unrelated problems — two different questions, two different tools.",
    "Fix: added _escape_like() + escape=\"\\\\\" on both ilike() calls in crud.py.",
    [("corrective/05_refactoring/sonarqube_result.png",
      "Figure 1.4 — SonarQube Cloud: first-ever project analysis (whole-project baseline)")],
    "This was the first-ever SonarQube analysis of the project (no prior baseline), so it reports "
    "whole-project health: 1 security issue, 13 reliability issues, 32 maintainability issues, 0% "
    "coverage (no test suite exists), and 27.7% duplication. None of this is attributable to the 3-line "
    "fix itself — it became the baseline every later cell compares against.")

add_table(doc, ["Input", "Before fix", "After fix"],
    [["search=\"_\"", "4/4 tasks (wrong)", "1/1 task (correct)"],
     ["search=\"%\"", "would also over-match", "0/0 tasks (correct)"]])
doc.add_page_break()

# ============================================================
# 2. ADAPTIVE
# ============================================================
doc.add_heading("5. Adaptive Maintenance", level=1)
add_para(doc,
    "Scenario: the project's Python runtime moved to Python 3.12+ (this environment runs 3.14.6), which "
    "deprecates datetime.datetime.utcnow() — used in 6 places across 3 files (auth.py, crud.py, "
    "models.py) — in favor of timezone-aware datetime.now(timezone.utc). A related SQLAlchemy 2.0 "
    "deprecation (declarative_base moved from sqlalchemy.ext.declarative to sqlalchemy.orm) was found and "
    "fixed alongside it, since it surfaced as a second blocking warning while verifying the same "
    "runtime-upgrade branch.", space_after=4)
add_rich_para(doc, [("Branch: ", True), ("adapt/datetime-utcnow-upgrade  ", False),
                     ("Fix commit: ", True), ("b177ef7  ", False),
                     ("Status: ", True), ("merged into main", False)], size=10, space_after=14)

add_cell_section(doc, "5.1", "Program Comprehension",
    "AST Explorer + written report",
    "No PySnooper here, because there is no wrong output to trace — the old code was functionally "
    "correct. The question is purely “which AST nodes tie this code to the old API,” which is "
    "exactly what AST Explorer answers.",
    None,
    [("adaptive/01_program_comprehension/ast_explorer_create_access_token.png",
      "Figure 2.1 — AST Explorer: create_access_token(), the two bare datetime.utcnow() calls")],
    "The AST shows exactly two bare Call nodes to datetime.utcnow() (lines 40 and 42), with no timezone "
    "object attached to either. A bonus finding surfaced while writing this report: because "
    "datetime.utcnow() returns a naive datetime, and PyJWT converts it to a Unix timestamp via "
    ".timestamp() (which reads a naive datetime as local time, not UTC), the old code had a latent "
    "timezone-correctness bug — verified directly on this machine as a 6-hour discrepancy between the "
    "naive-datetime timestamp and the real UTC epoch. This adaptation was not just silencing a warning, "
    "it fixed a real bug that would only show up on non-UTC servers.")

add_cell_section(doc, "5.2", "Change Management",
    "Git",
    "Same role as in the corrective branch — tracking the change and its justification.",
    None,
    None,
    "No requirements.txt version bump was needed — this is a runtime compatibility fix, not a dependency "
    "upgrade; the installed package versions already supported the new APIs used.")

add_cell_section(doc, "5.3", "Impact Analysis",
    "Joern + Neo4j",
    "The same CPG-query approach as the corrective branch, applied to a different question: scoping how "
    "widespread the deprecated-API usage is before touching any of it.",
    "# CPG built from the PRE-fix source, to scope the change rather than audit it after the fact\n"
    "cpg.call.code(\".*utcnow.*\").l   // every call site matching the old API",
    [("adaptive/03_impact_analysis/neo4j_impact_analysis.png",
      "Figure 2.2 — Neo4j Browser: the datetime.utcnow() adaptation surface")],
    "Confirmed exactly 3 files / 5 functions / 6 call sites — matching what was found by hand, nothing "
    "missed.")

add_cell_section(doc, "5.4", "Reverse Engineering",
    "Graphviz",
    "Same rationale as the corrective branch: turn the CPG's answer into a visual scope map.",
    None,
    [("adaptive/04_reverse_engineering/migration_surface.png",
      "Figure 2.3 — Graphviz: full migration surface vs. untouched parts of the project")],
    "Visualizes the full migration surface (the 5 affected functions) contrasted against untouched parts "
    "of the project (main.py endpoints, schemas.py, get_tasks()) — makes the “this touches exactly "
    "these 5 things and nothing else” claim visible at a glance.")

add_cell_section(doc, "5.5", "Refactoring",
    "SonarQube",
    "SonarQube checks that migrating 6 call sites did not introduce new bugs or smells project-wide.",
    "Fix: all 6 datetime.utcnow() call sites migrated to datetime.now(timezone.utc);\n"
    "database.py's declarative_base import corrected.",
    [("adaptive/05_refactoring/sonarqube_result.png",
      "Figure 2.4 — SonarQube Cloud: new-code quality gate for this branch")],
    "Both failed conditions on new code (0.0% coverage, 60.7% duplicated lines) were investigated and "
    "neither traces back to this adaptation. Coverage is the same pre-existing gap as the corrective "
    "baseline. The 60.7% duplication was traced to maintenance/'s own intentional before/after source "
    "snapshots fed to Joern being scanned as if they were production code — fixed by adding "
    "sonar.exclusions=maintenance/** to sonar-project.properties, which is the same fix that eventually "
    "drops duplication to 0.0% by the preventive branch. No new bugs or smells were reported in "
    "auth.py, crud.py, models.py, or database.py, the actual files this adaptation touched.")

add_table(doc, ["Condition (new code)", "Result", "Required"],
    [["Coverage", "0.0%", ">= 80.0%"],
     ["Duplicated Lines", "60.7%", "<= 3.0%"]])
doc.add_page_break()

# ============================================================
# 3. PREVENTIVE
# ============================================================
doc.add_heading("6. Preventive Maintenance", level=1)
add_para(doc,
    "Scenario: no active bug — this branch proactively looked for risky code before it caused one. "
    "static/app.js (653 lines, the largest and only untested/unstructured file in the project — no "
    "framework, no imports) had 4 functions (fetchTasks, handleTaskSubmit, toggleTaskStatus, deleteTask) "
    "each independently reimplementing the same fetch/401-check/error-handling boilerplate — exactly the "
    "kind of duplication that causes inconsistent fixes down the line.", space_after=4)
add_rich_para(doc, [("Branch: ", True), ("preventive/app-js-fetch-dedup  ", False),
                     ("Fix commit: ", True), ("a382f21  ", False),
                     ("Status: ", True), ("merged into main", False)], size=10, space_after=14)

add_cell_section(doc, "6.1", "Program Comprehension",
    "Graphviz + AST Explorer + written report",
    "Graphviz gives the project-wide structural overview needed to find a risky module in the first "
    "place — before knowing what is risky, AST Explorer cannot be targeted at anything specific. AST "
    "Explorer then confirms the specific duplication is real, not just eyeballed.",
    None,
    [("preventive/01_program_comprehension/project_structure.png",
      "Figure 3.1 — Graphviz: project-wide module dependency graph, app.js flagged"),
     ("preventive/01_program_comprehension/ast_explorer_toggleTaskStatus.png",
      "Figure 3.2 — AST Explorer: toggleTaskStatus()"),
     ("preventive/01_program_comprehension/ast_explorer_deleteTask.png",
      "Figure 3.3 — AST Explorer: deleteTask()")],
    "app.js stands out immediately in the project-wide graph (largest file, no tests, no framework). The "
    "two AST screenshots show toggleTaskStatus() and deleteTask() producing the identical "
    "TryStatement/IfStatement/CatchClause tree shape — a real structural duplicate, not just "
    "similar-looking code.")

add_cell_section(doc, "6.2", "Change Management",
    "Git",
    "Same role as the other branches.",
    None,
    None,
    "Per the plan's cell ordering, the actual code change was deliberately deferred to section 6.5, after "
    "Impact Analysis and clone detection quantified the duplication with tooling rather than just by eye.")

add_cell_section(doc, "6.3", "Impact Analysis",
    "SonarQube + Neo4j (via Joern)",
    "Joern's JavaScript frontend (jssrc2cpg) can quantify exactly how central each duplicated function "
    "is — call fan-out, who calls whom — turning “this looks duplicated” into “these 4 "
    "functions, confirmed, here is the call graph.”",
    "# CPG for app.js alone via jssrc2cpg\n"
    "cpg.call.name(\"handleLogout\").method.name.l.distinct   // every function using the 401-handling branch",
    [("preventive/03_impact_analysis/neo4j_impact_analysis.png",
      "Figure 3.4 — Neo4j Browser: app.js duplication cluster")],
    "The CPG query refined the comprehension finding from 3 to 4 duplicated functions — fetchTasks() "
    "itself turned out to share the identical shape too, something the manual reading had missed. This "
    "is exactly the value tooling adds over reading alone.")

add_cell_section(doc, "6.4", "Reverse Engineering",
    "CCFinderSW",
    "The plan calls for dedicated clone detection here, and CCFinderSW is a real, published, "
    "token-based clone detector (Kamiya et al.) — the right category of tool for “find duplicated "
    "logic blocks,” distinct from Joern's structural/call-graph analysis.",
    "CCFinderSW.bat D -d js_as_java -l java -o appjs_clones -ccfsw pair -t 30\n"
    "# .java extension trick used: this install has no native JavaScript ruleset;\n"
    "# Java's brace/semicolon/comment syntax tokenizes a .js file's shape closely enough to work.",
    [("preventive/04_reverse_engineering/ccfindersw_result.png",
      "Figure 3.5 — CCFinderSW terminal output")],
    "132 clone-pair entries found. The meaningful ones (a 4-way clique across fetchTasks / "
    "handleTaskSubmit / toggleTaskStatus / deleteTask) independently confirm — via pure token-based "
    "analysis, with zero knowledge of the AST/CPG findings — the exact same 4 functions. Three "
    "independent tools converged on the same answer. A bonus, explicitly out-of-scope finding also "
    "surfaced (handleLogin/handleRegister share structure too), flagged for a future pass rather than "
    "pulled into this branch.")

add_cell_section(doc, "6.5", "Refactoring",
    "SonarQube (verification) + CCFinderSW (re-run)",
    "The CCFinderSW re-run proves the targeted duplication is actually gone, not just “we added a "
    "helper and hope.” The SonarQube re-run proves nothing new broke project-wide.",
    "Fix: extracted a shared apiRequest() helper; rewired all 4 functions to use it.",
    [("preventive/05_refactoring/sonarqube_result.png",
      "Figure 3.6 — SonarQube Cloud: preventive branch result")],
    "fetchTasks and handleTaskSubmit now have zero clone matches with anything in the cluster. The one "
    "remaining match (toggleTaskStatus <-> deleteTask) is just the shared 3-line tail after calling the "
    "helper, not the risky duplicated auth/error logic that justified the branch — that logic now lives "
    "in exactly one place. Duplication dropping to 0.0% is partly this fix, partly the sonar.exclusions "
    "fix from section 5.5 finally taking full effect.")

add_table(doc, ["Metric", "Corrective baseline", "Preventive branch (after fix)"],
    [["Security", "1 open issue", "1 open issue"],
     ["Reliability", "13 open issues", "4 open issues"],
     ["Maintainability", "32 open issues", "16 open issues"],
     ["Duplications", "27.7%", "0.0%"],
     ["CCFinderSW pairwise matches (4 target fns)", "5 of 6 possible", "1 (trivial residual)"]])
doc.add_page_break()

# ============================================================
# 4. PERFECTIVE
# ============================================================
doc.add_heading("7. Perfective Maintenance", level=1)
add_para(doc,
    "Scenario: nothing is broken — crud.get_tasks() (and get_task()) return correct results for every "
    "filter today. But Task.user_id, .status, .priority, and .tag are all filtered on without a database "
    "index, so every call does a full table scan. Not a problem yet at the app's current scale, worth "
    "fixing proactively before it is. Reducing PBKDF2's iteration count on password hashing was "
    "considered and rejected as a “performance win” — that would be a security regression, not "
    "an optimization — so a different, safe target was chosen instead.", space_after=4)
add_rich_para(doc, [("Branch: ", True), ("perfective/get-tasks-query-indexing  ", False),
                     ("Fix commit: ", True), ("e4ba116  ", False),
                     ("Status: ", True), ("merged into main", False)], size=10, space_after=14)

add_cell_section(doc, "7.1", "Program Comprehension",
    "Viztracer + AST Explorer + written report",
    "Viztracer gives a call-timeline view of where time actually goes inside one get_tasks() call — "
    "useful for a first “what is even happening here” look. AST Explorer then structurally "
    "confirms exactly which parameters become WHERE-clause equality filters (the indexing candidates), "
    "separate from search's LIKE pattern which cannot benefit from a plain index the same way.",
    "# Seeded a realistic 20,000-row benchmark dataset (50 users x 400 tasks)\n"
    "# into an isolated SQLite file, traced one representative get_tasks() call with Viztracer.",
    [("perfective/01_program_comprehension/ast_explorer_get_tasks.png",
      "Figure 4.1 — AST Explorer: get_tasks() filter chain")],
    "sqlite3.Cursor.execute plus SQL compilation together accounted for the bulk of a 7.5 millisecond "
    "traced call — on an unindexed table. The AST confirmed structurally that user_id, status, priority, "
    "and tag all get plain equality filters (index candidates), while search gets an ilike() pattern "
    "(not the same kind of candidate).")

add_cell_section(doc, "7.2", "Change Management",
    "Git",
    "Same role as the other branches.",
    None,
    None,
    "The actual index-adding change was deferred to section 7.5, after Impact Analysis confirmed it was "
    "safe and the cProfile/Snakeviz pass pinpointed exactly which part of the call was the bottleneck.")

add_cell_section(doc, "7.3", "Impact Analysis",
    "Joern + Neo4j",
    "Same CPG-query approach as the corrective and adaptive branches, applied to confirm the enhancement "
    "cannot break any consumer.",
    "cpg.method.name(\"get_tasks\").caller.fullName.foreach(println)\n"
    "// plus every reference to models.Task.user_id / .status / .priority / .tag",
    [("perfective/03_impact_analysis/neo4j_impact_analysis.png",
      "Figure 4.2 — Neo4j Browser: get_tasks() / get_task() impact")],
    "get_tasks() still has exactly one caller (main.read_tasks) — unchanged from the corrective branch's "
    "finding. A bonus finding: get_task() (singular) also filters on user_id, so it gets the same "
    "speed-up for free. No semantic risk: indexes change lookup speed, never query results, so nothing "
    "could behave differently after this change.")

add_cell_section(doc, "7.4", "Reverse Engineering",
    "cProfile + Snakeviz",
    "Viztracer (section 7.1) shows a call-timeline overview; cProfile gives exact call-count and timing "
    "numbers needed to say precisely how much of the time is the database layer versus everything else. "
    "Snakeviz turns the resulting .prof file into an interactive flame graph.",
    "   ncalls  tottime  percall  cumtime  percall filename:lineno(function)\n"
    "       61    0.087    0.001    0.087    0.001 {method 'execute' of 'sqlite3.Cursor' objects}\n"
    "    12517    0.025    0.000    0.080    0.000 sqlalchemy/orm/loading.py:1068(_instance)\n"
    "    12517    0.020    0.000    0.020    0.000 sqlalchemy/orm/loading.py:1329(_populate_full)\n"
    "       60    0.018    0.000    0.018    0.000 {method 'fetchall' of 'sqlite3.Cursor' objects}",
    None,
    "sqlite3.Cursor.execute alone is 0.087s of the 0.301s total (29%) — the single largest self-time "
    "contributor by a wide margin, confirming with call-count precision that the bottleneck is the "
    "database layer, not get_tasks()'s own Python logic (0.001s total self time across all 60 calls). "
    "(Snakeviz's interactive flame-graph view was not screenshotted for this report — the text summary "
    "above captures the same finding.)")

add_cell_section(doc, "7.5", "Refactoring",
    "Loguru (verification) + cProfile/Snakeviz (re-check)",
    "Loguru confirms the migration logic itself behaves correctly (fires only when needed, silent when "
    "already applied) — a different kind of correctness check than “is it faster,” which is "
    "what the cProfile re-run answers.",
    "Fix: index=True on the 4 columns in models.py; a migration helper\n"
    "ensure_task_indexes() in main.py so an EXISTING todo.db gets retrofitted,\n"
    "not just fresh databases created by create_all().",
    None,
    "A real, measured, but honestly modest improvement — not a dramatic speedup, and reported as such "
    "rather than oversold. SQLite was already reasonably fast at 20,000 rows; the gap widens as the "
    "table grows further, which is the actual point of doing this proactively. Loguru confirmed the "
    "migration fires exactly 4 times on a legacy unindexed database and zero times on a re-run — "
    "correct and idempotent.")

add_table(doc, ["Metric", "Before", "After", "Change"],
    [["Total wall time (60 calls)", "0.301s", "0.245s", "-19%"],
     ["sqlite3.Cursor.execute self time", "0.087s", "0.063s", "-28%"],
     ["get_tasks() self time", "0.001s", "~0.001s", "unchanged (expected)"]])
doc.add_page_break()

# ---------------- Consolidated results ----------------
doc.add_heading("8. Consolidated Results", level=1)
add_table(doc,
    ["Branch", "Commit", "Real finding", "Verified fix"],
    [
        ["fix/search-like-wildcard-escape", "f4dc3cf",
         "LIKE-wildcard injection in search (wrong results, not a crash)",
         "PySnooper 4/4 -> 1/1 correct; SonarQube baseline established"],
        ["adapt/datetime-utcnow-upgrade", "b177ef7",
         "Deprecated datetime.utcnow() (6 sites) + latent JWT-expiry timezone bug (6h)",
         "0 targeted deprecation warnings remain; both issues fixed together"],
        ["preventive/app-js-fetch-dedup", "a382f21",
         "4-function fetch/401/error-handling duplication in app.js",
         "CCFinderSW 5/6 -> 1 pairwise matches; SonarQube duplication 27.7% -> 0.0%"],
        ["perfective/get-tasks-query-indexing", "e4ba116",
         "Unindexed columns on every get_tasks()/get_task() call",
         "cProfile -19% wall time, -28% on the bottleneck"],
    ])

doc.add_heading("9. Reflections", level=1)
add_para(doc,
    "Working through all four categories back to back on the same real codebase made the differences "
    "between them concrete rather than theoretical. Corrective and adaptive maintenance both start from "
    "“something is provably wrong” — a wrong output, a deprecation warning — and the tools used "
    "are about proving the fix is complete and safe. Preventive and perfective maintenance both start "
    "from “nothing is wrong yet,” and the tools used are about finding something worth fixing in "
    "the first place (duplication, an unindexed hot path) and then honestly measuring whether the fix was "
    "actually worth it.")
add_para(doc,
    "The branches were not run in isolation. The corrective fix's SonarQube run doubled as the "
    "project-wide baseline; the adaptive branch's duplication spike was explained and its root cause "
    "fixed via a scanner-configuration change; and the preventive branch's 0.0% duplication result then "
    "depended on that earlier fix taking effect. This is a reasonably faithful picture of how real "
    "maintenance history works — later work keeps building on and correcting earlier work, rather than "
    "each change happening in a vacuum.")

doc.save(os.path.join(HERE, "FINAL_REPORT.docx"))
print("Saved:", os.path.join(HERE, "FINAL_REPORT.docx"))
